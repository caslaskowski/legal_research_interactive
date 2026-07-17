#!/usr/bin/env python3
"""Generate a Word companion-notes document for every module that has companion data.

Two places to put companion content, one output convention:

  1. A module's own  <module>/data/companion.json           (the ch-1 orientation modules)
  2. A root-level    data/companions/<anything>.json        (modules without a data/ dir;
                                                             the JSON names its module via "module")

Either way the schema is the same:

  {
    "module":  "ch-2/using-databases",       # only needed for files under data/companions/
    "title":   "Using Databases — Companion Notes",
    "subtitle": "Ch. 2 · Introduction to Legal Research Platforms",   # optional
    "intro":   "One short paragraph telling the student how to use the sheet.",
    "sections": [
      { "heading": "SECTION NAME",
        "fields": [ { "label": "Prompt the student fills in", "lines": 2 } ] }
    ],
    "footer":  "Closing line."               # optional
  }

Output: <module>/companion-notes.docx — a static file, committed to the repo, so the
site keeps its zero-external-requests posture and students download it like any asset.

Run:  python3 build_companions.py          (requires:  pip install python-docx)
"""
import glob
import json
import os
import sys

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    sys.exit("python-docx is not installed. Run:  pip install python-docx")

HERE = os.path.dirname(os.path.abspath(__file__))

INK = RGBColor(0x26, 0x26, 0x22)       # site --ink
FLAME = RGBColor(0xA7, 0x2D, 0x2A)     # site --flame
FAINT = RGBColor(0x6B, 0x6A, 0x62)     # site --ink-faint
RULE = "B6B6B6"                        # site --rule-strong
SANS = "Arial"                          # stand-ins for Archivo / Newsreader that every
SERIF = "Georgia"                       # copy of Word can render without embedding fonts


def _set_font(run, name, size, color=INK, bold=False, italic=False, caps=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    if caps:
        run.font.all_caps = True
    # ensure the east-asian font slot matches, or Word substitutes on some systems
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def _bottom_border(paragraph, color=RULE, size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pPr.makeelement(qn("w:bottom"), {})
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _writing_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    _bottom_border(p)
    return p


def build_docx(data, out_path):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)      # US Letter
    section.page_height = Inches(11)
    for side in ("left_margin", "right_margin"):
        setattr(section, side, Inches(1))
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

    # eyebrow
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("The Question Method of Legal Research \u00b7 Companion Notes")
    _set_font(r, SANS, 8.5, FLAME, bold=True, caps=True)

    # title
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(data["title"])
    _set_font(r, SANS, 20, INK, bold=True)

    # subtitle (chapter line)
    if data.get("subtitle"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(data["subtitle"])
        _set_font(r, SANS, 10, FAINT)

    # name / date line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Name: " + "\u2002" * 24 + "Date: " + "\u2002" * 12)
    _set_font(r, SANS, 9.5, FAINT)
    _bottom_border(p)

    # intro
    if data.get("intro"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(data["intro"])
        _set_font(r, SERIF, 10.5, INK, italic=True)

    for sec in data.get("sections", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(sec["heading"])
        _set_font(r, SANS, 11, FLAME, bold=True, caps=True)
        _bottom_border(p, color="A72D2A", size=10)

        for field in sec.get("fields", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(field["label"])
            _set_font(r, SANS, 10, INK, bold=True)
            for _ in range(int(field.get("lines", 2))):
                _writing_line(doc)

    if data.get("footer"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(data["footer"])
        _set_font(r, SERIF, 9.5, FAINT, italic=True)

    doc.save(out_path)


def collect():
    """Yield (companion_data, module_dir) for every companion source in the repo."""
    for path in sorted(glob.glob(os.path.join(HERE, "ch-*", "*", "data", "companion.json"))):
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        yield data, os.path.dirname(os.path.dirname(path))
    for path in sorted(glob.glob(os.path.join(HERE, "data", "companions", "*.json"))):
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        module = data.get("module")
        if not module:
            print("  ! %s has no \"module\" field — skipped" % os.path.relpath(path, HERE))
            continue
        mod_dir = os.path.join(HERE, module)
        if not os.path.isdir(mod_dir):
            print("  ! %s points at missing module %s — skipped" % (os.path.relpath(path, HERE), module))
            continue
        yield data, mod_dir


def main():
    count = 0
    for data, mod_dir in collect():
        out = os.path.join(mod_dir, "companion-notes.docx")
        build_docx(data, out)
        print("  wrote", os.path.relpath(out, HERE))
        count += 1
    print("companion notes: %d document(s) generated" % count)


if __name__ == "__main__":
    main()
